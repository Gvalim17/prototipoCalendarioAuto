"""Geração de PDF/Word do Plano de Trabalho Docente (PTD) para entrega à
coordenação. Usa bibliotecas puramente Python (sem dependências de sistema
como Cairo/Pango), compatíveis com o deploy no Render."""

import io
from xml.sax.saxutils import escape

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


SECTIONS = [
    ("ementa", "Ementa"),
    ("objetivos", "Objetivos"),
    ("conteudo_programatico", "Conteúdo Programático"),
    ("metodologia", "Metodologia"),
    ("recursos_didaticos", "Recursos Didáticos"),
    ("criterios_avaliacao", "Critérios de Avaliação"),
    ("bibliografia", "Bibliografia"),
    ("notes", "Observações"),
]


def _field(plan, key: str) -> str:
    if plan is None:
        return ""
    value = getattr(plan, key, None) if not isinstance(plan, dict) else plan.get(key)
    return (value or "").strip()


def render_lesson_plan_docx(discipline_name: str, plan) -> io.BytesIO:
    document = Document()
    title = document.add_heading("Plano de Trabalho Docente (PTD)", level=1)
    title.runs[0].font.size = Pt(18)
    document.add_heading(discipline_name, level=2)

    for key, label in SECTIONS:
        content = _field(plan, key)
        if not content:
            continue
        document.add_heading(label, level=3)
        for paragraph in content.split("\n"):
            document.add_paragraph(paragraph)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def render_lesson_plan_pdf(discipline_name: str, plan) -> io.BytesIO:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm)
    styles = getSampleStyleSheet()
    story = [
        Paragraph("Plano de Trabalho Docente (PTD)", styles["Title"]),
        Paragraph(escape(discipline_name), styles["Heading2"]),
        Spacer(1, 0.5 * cm),
    ]
    for key, label in SECTIONS:
        content = _field(plan, key)
        if not content:
            continue
        story.append(Paragraph(escape(label), styles["Heading3"]))
        for paragraph in content.split("\n"):
            if paragraph.strip():
                # Paragraph do reportlab interpreta uma mini-sintaxe tipo XML
                # (<font>, <b>...). Sem escapar, texto digitado pelo professor
                # (ex.: "<font color=\"x\">") derruba a geração do PDF com uma
                # exceção — inclusive quando é o admin quem exporta.
                story.append(Paragraph(escape(paragraph), styles["BodyText"]))
        story.append(Spacer(1, 0.3 * cm))

    doc.build(story)
    buffer.seek(0)
    return buffer


def _script_field(script, key: str) -> str:
    if script is None:
        return ""
    value = getattr(script, key, None) if not isinstance(script, dict) else script.get(key)
    return (value or "").strip()


def render_lesson_script_docx(discipline_name: str, class_date: str, script) -> io.BytesIO:
    document = Document()
    title = document.add_heading("Roteiro de Aula", level=1)
    title.runs[0].font.size = Pt(18)
    document.add_heading(f"{discipline_name} · {class_date}", level=2)

    topic = _script_field(script, "topic")
    if topic:
        document.add_heading("Tema do dia", level=3)
        document.add_paragraph(topic)

    content = _script_field(script, "content")
    if content:
        document.add_heading("Roteiro / atividades", level=3)
        for paragraph in content.split("\n"):
            document.add_paragraph(paragraph)

    attachments = script.attachments if not isinstance(script, dict) else script.get("attachments", [])
    if attachments:
        document.add_heading("Materiais anexados", level=3)
        for att in attachments:
            document.add_paragraph(att.filename, style="List Bullet")

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def render_lesson_script_pdf(discipline_name: str, class_date: str, script) -> io.BytesIO:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm)
    styles = getSampleStyleSheet()
    story = [
        Paragraph("Roteiro de Aula", styles["Title"]),
        Paragraph(escape(f"{discipline_name} · {class_date}"), styles["Heading2"]),
        Spacer(1, 0.5 * cm),
    ]

    topic = _script_field(script, "topic")
    if topic:
        story.append(Paragraph("Tema do dia", styles["Heading3"]))
        story.append(Paragraph(escape(topic), styles["BodyText"]))
        story.append(Spacer(1, 0.3 * cm))

    content = _script_field(script, "content")
    if content:
        story.append(Paragraph("Roteiro / atividades", styles["Heading3"]))
        for paragraph in content.split("\n"):
            if paragraph.strip():
                story.append(Paragraph(escape(paragraph), styles["BodyText"]))
        story.append(Spacer(1, 0.3 * cm))

    attachments = script.attachments if not isinstance(script, dict) else script.get("attachments", [])
    if attachments:
        story.append(Paragraph("Materiais anexados", styles["Heading3"]))
        for att in attachments:
            story.append(Paragraph(f"• {escape(att.filename)}", styles["BodyText"]))

    doc.build(story)
    buffer.seek(0)
    return buffer
